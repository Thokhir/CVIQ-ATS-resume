"""
exporter.py — Template-aware DOCX / TXT export for ATS Resume Maker Pro.

Each template applies its own:
  - fonts, accent colours, section heading style
  - section ORDER (e.g. education-first for freshers)
  - margins, spacing

DOCX is built with python-docx.
"""
import io
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Template definitions ──────────────────────────────────────────────────────

TEMPLATES: dict[str, dict] = {
    "classic_professional": {
        "name": "Classic Professional",
        "desc": "Universal ATS-safe • Calibri • Dark navy headings",
        "category": "all", "free": True,
        "font": "Calibri", "font_size": 11,
        "heading_font": "Calibri", "heading_size": 12,
        "name_size": 18,
        "accent": RGBColor(0x1F, 0x49, 0x7D),   # dark navy
        "text_color": RGBColor(0x00, 0x00, 0x00),
        "heading_style": "underline",             # underline or border
        "margins": (2.0, 1.5, 2.0, 1.5),         # top, right, bot, left (cm)
        "line_spacing": 1.15,
        "section_order": ["summary", "experience", "education", "skills",
                          "certifications", "publications", "patents",
                          "awards", "projects", "languages", "activities"],
    },
    "modern_clean": {
        "name": "Modern Clean",
        "desc": "Blue accents • Contemporary • Ideal for tech roles",
        "category": "tech", "free": True,
        "font": "Calibri", "font_size": 11,
        "heading_font": "Calibri", "heading_size": 12,
        "name_size": 20,
        "accent": RGBColor(0x2E, 0x75, 0xB6),   # modern blue
        "text_color": RGBColor(0x1A, 0x1A, 0x2E),
        "heading_style": "border",
        "margins": (1.8, 1.5, 1.8, 1.5),
        "line_spacing": 1.15,
        "section_order": ["summary", "skills", "experience", "education",
                          "certifications", "projects", "awards", "languages"],
    },
    "fresh_graduate": {
        "name": "Fresh Graduate",
        "desc": "Education-first • Projects & Internships highlighted",
        "category": "fresher", "free": True,
        "font": "Calibri", "font_size": 11,
        "heading_font": "Calibri", "heading_size": 12,
        "name_size": 18,
        "accent": RGBColor(0x05, 0x78, 0x5B),   # emerald green
        "text_color": RGBColor(0x11, 0x18, 0x27),
        "heading_style": "underline",
        "margins": (2.0, 1.5, 2.0, 1.5),
        "line_spacing": 1.2,
        "section_order": ["summary", "education", "skills", "projects",
                          "experience", "certifications", "awards",
                          "activities", "languages"],
    },
    "executive": {
        "name": "Executive",
        "desc": "Georgia serif • Senior & leadership roles",
        "category": "business", "free": False,
        "font": "Georgia", "font_size": 11,
        "heading_font": "Georgia", "heading_size": 12,
        "name_size": 18,
        "accent": RGBColor(0x1A, 0x1A, 0x2E),
        "text_color": RGBColor(0x11, 0x11, 0x11),
        "heading_style": "border",
        "margins": (2.2, 2.0, 2.2, 2.0),
        "line_spacing": 1.2,
        "section_order": ["summary", "credentials", "experience", "education",
                          "skills", "certifications", "publications", "patents",
                          "awards", "languages"],
    },
    "academic_research": {
        "name": "Academic / Research",
        "desc": "Times New Roman • Publications prominent • PhD/PostDoc",
        "category": "academic", "free": False,
        "font": "Times New Roman", "font_size": 11,
        "heading_font": "Times New Roman", "heading_size": 12,
        "name_size": 18,
        "accent": RGBColor(0x44, 0x72, 0xC4),
        "text_color": RGBColor(0x00, 0x00, 0x00),
        "heading_style": "underline",
        "margins": (2.0, 2.0, 2.0, 2.0),
        "line_spacing": 1.15,
        "section_order": ["summary", "education", "publications", "patents",
                          "experience", "skills", "certifications", "awards",
                          "projects", "languages"],
    },
    "pharma_cro": {
        "name": "Pharma / CRO",
        "desc": "Purple accents • Life sciences & drug discovery",
        "category": "science", "free": False,
        "font": "Calibri", "font_size": 11,
        "heading_font": "Calibri", "heading_size": 12,
        "name_size": 18,
        "accent": RGBColor(0x7B, 0x1F, 0xA2),
        "text_color": RGBColor(0x11, 0x11, 0x11),
        "heading_style": "border",
        "margins": (2.0, 1.5, 2.0, 1.5),
        "line_spacing": 1.15,
        "section_order": ["summary", "credentials", "experience", "education",
                          "skills", "publications", "patents", "certifications",
                          "awards", "projects", "languages"],
    },
    "data_science": {
        "name": "Data Science / AI",
        "desc": "Teal accents • ML Engineer • Data Scientist",
        "category": "tech", "free": False,
        "font": "Arial", "font_size": 11,
        "heading_font": "Arial", "heading_size": 12,
        "name_size": 18,
        "accent": RGBColor(0x00, 0x89, 0x7B),
        "text_color": RGBColor(0x11, 0x11, 0x11),
        "heading_style": "border",
        "margins": (1.8, 1.5, 1.8, 1.5),
        "line_spacing": 1.15,
        "section_order": ["summary", "skills", "experience", "projects",
                          "education", "certifications", "publications",
                          "awards", "languages"],
    },
    "healthcare": {
        "name": "Healthcare",
        "desc": "Red accents • Medical & clinical roles",
        "category": "science", "free": False,
        "font": "Calibri", "font_size": 11,
        "heading_font": "Calibri", "heading_size": 12,
        "name_size": 18,
        "accent": RGBColor(0xDC, 0x26, 0x26),
        "text_color": RGBColor(0x11, 0x11, 0x11),
        "heading_style": "underline",
        "margins": (2.0, 1.5, 2.0, 1.5),
        "line_spacing": 1.2,
        "section_order": ["summary", "experience", "education", "certifications",
                          "skills", "awards", "languages"],
    },
}


def get_all_templates() -> dict:
    return TEMPLATES


def get_template_names() -> list[str]:
    return list(TEMPLATES.keys())


def _resolve_template(template_id: str) -> dict:
    """
    Resolve a template id to a build spec from EITHER the legacy 8 exporter
    templates OR the 21 gallery templates (template_manager.BUILTIN_TEMPLATES).
    This single resolver is why the template a user selects in the gallery is
    the template that actually gets applied to the exported DOCX.
    """
    tid = (template_id or "").strip()
    # Word-file templates are handled upstream; fall back to a safe default here.
    if tid.startswith("word_"):
        return TEMPLATES["classic_professional"]
    if tid in TEMPLATES:
        return TEMPLATES[tid]
    try:
        from utils.template_manager import BUILTIN_TEMPLATES, get_export_spec
        if tid in BUILTIN_TEMPLATES:
            return get_export_spec(tid)
    except Exception:
        pass
    return TEMPLATES["classic_professional"]


def template_exists(template_id: str) -> bool:
    tid = (template_id or "").strip()
    if tid in TEMPLATES:
        return True
    try:
        from utils.template_manager import BUILTIN_TEMPLATES
        return tid in BUILTIN_TEMPLATES
    except Exception:
        return False


# ── DOCX Builder ──────────────────────────────────────────────────────────────

def _rgb(r: RGBColor):
    return r


def _set_run_font(run, font_name: str, size: int, bold=False, color: RGBColor = None, italic=False):
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def _add_heading(doc: Document, text: str, tpl: dict):
    style = tpl["heading_style"]
    accent = tpl["accent"]
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(10)
    para.paragraph_format.space_after = Pt(2)

    run = para.add_run(text.upper())
    _set_run_font(run, tpl["heading_font"], tpl["heading_size"], bold=True, color=accent)

    if style == "border":
        _add_bottom_border(para, accent)
    else:
        # underline the run
        run.underline = True
    return para


def _add_bottom_border(para, color: RGBColor):
    """Add a bottom border to a paragraph in the template accent colour."""
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    hex_color = str(color)  # RGBColor.__str__ returns hex like '1F497D'
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), hex_color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_name_contact(doc: Document, data: dict, tpl: dict):
    pi = data.get("personal_info", {})
    name = pi.get("name") or data.get("full_name") or ""
    email = pi.get("email", "")
    phone = pi.get("phone", "")
    location = pi.get("location", "")
    linkedin = pi.get("linkedin", "")
    github = pi.get("github", "")
    website = pi.get("website", "")
    accent = tpl["accent"]

    # Name
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_para.paragraph_format.space_after = Pt(2)
    name_run = name_para.add_run(name)
    _set_run_font(name_run, tpl["heading_font"], tpl["name_size"], bold=True, color=accent)

    # Contact line
    contact_parts = [p for p in [email, phone, location, linkedin, github, website] if p]
    if contact_parts:
        c_para = doc.add_paragraph()
        c_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        c_para.paragraph_format.space_after = Pt(6)
        c_run = c_para.add_run("  |  ".join(contact_parts))
        _set_run_font(c_run, tpl["font"], 9, color=RGBColor(0x44, 0x44, 0x44))


def _section_text_to_docx(doc: Document, section_text: str, tpl: dict):
    """Parse section body lines and add them to the document."""
    BULLET_RE = re.compile(r"^[•\-\*–◦·]\s*(.+)")
    lines = section_text.split("\n")
    fn = tpl["font"]
    fs = tpl["font_size"]
    tc = tpl["text_color"]

    for line in lines:
        s = line.strip()
        if not s:
            continue
        m = BULLET_RE.match(s)
        if m:
            para = doc.add_paragraph(style="List Bullet")
            para.paragraph_format.left_indent = Cm(0.5)
            para.paragraph_format.space_before = Pt(1)
            para.paragraph_format.space_after = Pt(1)
            run = para.add_run(m.group(1))
            _set_run_font(run, fn, fs, color=tc)
        else:
            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(1)
            para.paragraph_format.space_after = Pt(1)
            run = para.add_run(s)
            _set_run_font(run, fn, fs, color=tc)


def _build_docx_from_resume_data(data: dict, template_id: str) -> bytes:
    """
    Build a DOCX from structured resume_data dict (from the Builder form).
    Data keys: personal_info, summary, experience[], education[], skills[],
    certifications[], publications[], projects[], awards[], languages, activities
    """
    tpl = _resolve_template(template_id)
    doc = Document()

    # Page margins
    t, r, b, l = tpl["margins"]
    for sec in doc.sections:
        sec.top_margin = Cm(t)
        sec.right_margin = Cm(r)
        sec.bottom_margin = Cm(b)
        sec.left_margin = Cm(l)

    # Header
    _add_name_contact(doc, data, tpl)

    section_order = tpl["section_order"]

    def _add_sec(key: str):
        if key == "summary":
            summary = data.get("summary", "").strip()
            if not summary:
                return
            _add_heading(doc, "Professional Summary", tpl)
            para = doc.add_paragraph()
            run = para.add_run(summary)
            _set_run_font(run, tpl["font"], tpl["font_size"], color=tpl["text_color"])

        elif key == "experience":
            experiences = data.get("experience", [])
            if not any(e.get("company") or e.get("position") for e in experiences):
                return
            _add_heading(doc, "Work Experience", tpl)
            for exp in experiences:
                pos = exp.get("position", "").strip()
                company = exp.get("company", "").strip()
                dates = exp.get("dates", "").strip()
                location = exp.get("location", "").strip()
                desc = exp.get("description", "").strip()
                if not pos and not company:
                    continue
                # Job title | Company line
                title_line = " | ".join(filter(None, [pos, company]))
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(6)
                r1 = p.add_run(title_line)
                _set_run_font(r1, tpl["font"], tpl["font_size"], bold=True, color=tpl["accent"])
                # Dates | Location
                meta = " | ".join(filter(None, [dates, location]))
                if meta:
                    p2 = doc.add_paragraph()
                    p2.paragraph_format.space_before = Pt(0)
                    p2.paragraph_format.space_after = Pt(1)
                    r2 = p2.add_run(meta)
                    _set_run_font(r2, tpl["font"], 9, italic=True,
                                  color=RGBColor(0x55, 0x55, 0x55))
                # Bullets
                if desc:
                    for line in desc.split("\n"):
                        s = line.strip()
                        if s:
                            bm = re.match(r"^[•\-\*–◦·]\s*(.+)", s)
                            bp = doc.add_paragraph(style="List Bullet")
                            bp.paragraph_format.left_indent = Cm(0.5)
                            bp.paragraph_format.space_before = Pt(1)
                            bp.paragraph_format.space_after = Pt(1)
                            run = bp.add_run(bm.group(1) if bm else s)
                            _set_run_font(run, tpl["font"], tpl["font_size"], color=tpl["text_color"])

        elif key == "education":
            educations = data.get("education", [])
            if not any(e.get("degree") or e.get("school") for e in educations):
                return
            _add_heading(doc, "Education", tpl)
            for edu in educations:
                degree = edu.get("degree", "").strip()
                school = edu.get("school", "").strip()
                year = str(edu.get("year", "")).strip()
                gpa = edu.get("gpa", "").strip()
                honors = edu.get("honors", "").strip()
                if not degree and not school:
                    continue
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(5)
                r1 = p.add_run(degree or school)
                _set_run_font(r1, tpl["font"], tpl["font_size"], bold=True, color=tpl["accent"])
                if degree and school:
                    r2 = p.add_run(f" — {school}")
                    _set_run_font(r2, tpl["font"], tpl["font_size"], color=tpl["text_color"])
                meta = " | ".join(filter(None, [year, gpa, honors]))
                if meta:
                    p2 = doc.add_paragraph()
                    p2.paragraph_format.space_before = Pt(0)
                    r2 = p2.add_run(meta)
                    _set_run_font(r2, tpl["font"], 9, italic=True,
                                  color=RGBColor(0x55, 0x55, 0x55))

        elif key == "skills":
            skills_raw = data.get("skills", "")
            if isinstance(skills_raw, list):
                skills_text = ", ".join(str(s).strip() for s in skills_raw if s)
            else:
                skills_text = str(skills_raw).strip()
            # Also collect categorised skills
            skill_cats = data.get("skill_categories", {})
            if not skills_text and not skill_cats:
                return
            _add_heading(doc, "Skills", tpl)
            if skill_cats:
                for cat, items in skill_cats.items():
                    if not items:
                        continue
                    p = doc.add_paragraph()
                    r1 = p.add_run(f"{cat}: ")
                    _set_run_font(r1, tpl["font"], tpl["font_size"], bold=True, color=tpl["text_color"])
                    r2 = p.add_run(", ".join(items) if isinstance(items, list) else str(items))
                    _set_run_font(r2, tpl["font"], tpl["font_size"], color=tpl["text_color"])
            elif skills_text:
                para = doc.add_paragraph()
                run = para.add_run(skills_text)
                _set_run_font(run, tpl["font"], tpl["font_size"], color=tpl["text_color"])

        elif key == "projects":
            projects = data.get("projects", [])
            if not any(p.get("title") for p in projects):
                return
            _add_heading(doc, "Projects", tpl)
            for proj in projects:
                title = proj.get("title", "").strip()
                tech = proj.get("tech", "").strip()
                link = proj.get("link", "").strip()
                desc = proj.get("description", "").strip()
                if not title:
                    continue
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(5)
                r1 = p.add_run(title)
                _set_run_font(r1, tpl["font"], tpl["font_size"], bold=True, color=tpl["accent"])
                if tech:
                    r2 = p.add_run(f" | {tech}")
                    _set_run_font(r2, tpl["font"], 9, color=RGBColor(0x55, 0x55, 0x55))
                if link:
                    r3 = p.add_run(f" | {link}")
                    _set_run_font(r3, tpl["font"], 9, italic=True, color=tpl["accent"])
                if desc:
                    for line in desc.split("\n"):
                        s = line.strip()
                        if s:
                            bp = doc.add_paragraph(style="List Bullet")
                            bp.paragraph_format.left_indent = Cm(0.5)
                            run = bp.add_run(s)
                            _set_run_font(run, tpl["font"], tpl["font_size"], color=tpl["text_color"])

        elif key == "certifications":
            certs = data.get("certifications", [])
            cert_str = data.get("certifications_text", "")
            if isinstance(certs, list):
                items = [c for c in certs if c.strip()]
            else:
                items = [c.strip() for c in str(certs).split("\n") if c.strip()]
            if cert_str:
                items += [c.strip() for c in cert_str.split("\n") if c.strip()]
            if not items:
                return
            _add_heading(doc, "Certifications", tpl)
            for item in items:
                bp = doc.add_paragraph(style="List Bullet")
                bp.paragraph_format.left_indent = Cm(0.5)
                run = bp.add_run(item)
                _set_run_font(run, tpl["font"], tpl["font_size"], color=tpl["text_color"])

        elif key == "publications":
            pubs = data.get("publications", [])
            pub_str = data.get("publications_text", "")
            items = (pubs if isinstance(pubs, list) else []) + [pub_str] if pub_str else pubs
            items = [str(i).strip() for i in items if str(i).strip()]
            if not items:
                return
            _add_heading(doc, "Publications", tpl)
            for i, item in enumerate(items, 1):
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(2)
                run = p.add_run(f"{i}. {item}")
                _set_run_font(run, tpl["font"], tpl["font_size"], color=tpl["text_color"])

        elif key == "patents":
            pat_str = data.get("patents_text", "").strip()
            patents = data.get("patents", [])
            items = [str(i).strip() for i in patents if str(i).strip()]
            if pat_str:
                items += [s.strip() for s in pat_str.split("\n") if s.strip()]
            if not items:
                return
            _add_heading(doc, "Patents", tpl)
            for item in items:
                p = doc.add_paragraph()
                run = p.add_run(item)
                _set_run_font(run, tpl["font"], tpl["font_size"], color=tpl["text_color"])

        elif key == "awards":
            awards = data.get("awards", [])
            aw_str = data.get("awards_text", "")
            items = (awards if isinstance(awards, list) else []) + ([aw_str] if aw_str else [])
            items = [str(i).strip() for i in items if str(i).strip()]
            if not items:
                return
            _add_heading(doc, "Awards & Honors", tpl)
            for item in items:
                bp = doc.add_paragraph(style="List Bullet")
                run = bp.add_run(item)
                _set_run_font(run, tpl["font"], tpl["font_size"], color=tpl["text_color"])

        elif key == "languages":
            lang = data.get("languages", "").strip()
            if not lang:
                return
            _add_heading(doc, "Languages", tpl)
            p = doc.add_paragraph()
            run = p.add_run(lang)
            _set_run_font(run, tpl["font"], tpl["font_size"], color=tpl["text_color"])

        elif key == "activities":
            act = data.get("activities", "").strip()
            if not act:
                return
            _add_heading(doc, "Activities & Extracurriculars", tpl)
            for line in act.split("\n"):
                s = line.strip()
                if s:
                    bp = doc.add_paragraph(style="List Bullet")
                    run = bp.add_run(s)
                    _set_run_font(run, tpl["font"], tpl["font_size"], color=tpl["text_color"])

        elif key == "credentials":
            cred = data.get("credentials", "").strip()
            if not cred:
                return
            _add_heading(doc, "Key Credentials", tpl)
            for line in cred.split("\n"):
                s = line.strip()
                if s:
                    bp = doc.add_paragraph(style="List Bullet")
                    run = bp.add_run(s)
                    _set_run_font(run, tpl["font"], tpl["font_size"], color=tpl["text_color"])

    for key in section_order:
        _add_sec(key)

    # Save to bytes
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _build_docx_from_text(text: str, template_id: str,
                          parsed_sections: dict = None) -> bytes:
    """
    Build properly formatted DOCX from optimizer output text.
    Uses pre-parsed sections dict (parsed_sections) when provided to avoid
    losing sections due to PDF page-break artifacts.
    All 4 jobs (Novoridge, SRKREC, Eesavyasa, CSIR-IICT) are preserved.
    """
    from utils.resume_processor import parse_resume_text, _clean_pdf_text

    tpl = _resolve_template(template_id)
    doc = Document()

    for sec in doc.sections:
        t, r, b, l = tpl["margins"]
        sec.top_margin    = Cm(t)
        sec.right_margin  = Cm(r)
        sec.bottom_margin = Cm(b)
        sec.left_margin   = Cm(l)

    # Use pre-parsed sections if available (most accurate)
    if parsed_sections and parsed_sections.get("_sections"):
        parsed = parsed_sections
        # But update the raw text to use the optimized version
        parsed = dict(parsed)
        # Rebuild _sections from optimized text so changes are included
        opt_parsed = parse_resume_text(text)
        # Merge: use optimized sections for summary/skills/experience, original for rest
        merged_sections = dict(parsed["_sections"])
        for key in ("summary", "skills", "experience"):
            if key in opt_parsed["_sections"]:
                merged_sections[key] = opt_parsed["_sections"][key]
        parsed["_sections"] = merged_sections
    else:
        parsed = parse_resume_text(_clean_pdf_text(text) if text else text)

    pi    = parsed.get("personal_info", {})
    order = parsed.get("_order", [])
    secs  = parsed.get("_sections", {})
    fn  = tpl["font"]
    hfn = tpl["heading_font"]
    fs  = tpl["font_size"]
    hs  = tpl["heading_size"]
    ns  = tpl["name_size"]
    accent = tpl["accent"]
    tc = RGBColor(0x22, 0x22, 0x22)

    # ── Name
    name = pi.get("name","").strip()
    if not name:
        for line in secs.get("_header_","").split("\n"):
            if line.strip():
                name = line.strip()
                break
    if name:
        p  = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        nr = p.add_run(name)
        _set_run_font(nr, hfn, ns, bold=True, color=accent)

    # ── Subtitle / tag line from header block
    hdr_text  = secs.get("_header_","")
    sub_re    = re.compile(
        r"(Research Scientist|Data Scientist|Scientist|Engineer|Analyst|Director|"
        r"Manager|Consultant|Researcher|Biologist|Chemist|Developer|Lead|Principal)"
        r"[^\n]{5,100}", re.I
    )
    sm = sub_re.search(hdr_text)
    if sm:
        p  = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        sr = p.add_run(sm.group(0).strip())
        _set_run_font(sr, fn, fs - 1, italic=True, color=accent)

    # ── Contact line
    contact_parts = [v for k,v in pi.items()
                     if k not in ("name",) and v and len(v.strip()) > 3 and len(v) < 100]
    if contact_parts:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(6)
        cr = cp.add_run("  |  ".join(contact_parts[:5]))
        _set_run_font(cr, fn, 9, color=RGBColor(0x44, 0x44, 0x44))

    # ── Section display names
    SEC_NAMES = {
        "summary":        "PROFESSIONAL SUMMARY",
        "experience":     "WORK EXPERIENCE",
        "education":      "EDUCATION",
        "skills":         "TECHNICAL SKILLS",
        "certifications": "CERTIFICATIONS",
        "publications":   "PUBLICATIONS",
        "patents":        "PATENTS",
        "awards":         "AWARDS & HONORS",
        "projects":       "PROJECTS",
        "languages":      "LANGUAGES",
        "credentials":    "KEY CREDENTIALS",
        "activities":     "EXTRACURRICULAR ACTIVITIES",
    }

    # Render order: template order first, then any extra sections in doc order
    tpl_order   = tpl.get("section_order", ["summary","experience","education",
                            "skills","certifications","publications","patents","awards"])
    doc_order   = [k for k in order if k != "_header_"]
    rendered    = set()
    render_q    = []
    for k in tpl_order:
        if k in doc_order and k not in rendered:
            render_q.append(k); rendered.add(k)
    for k in doc_order:
        if k not in rendered:
            render_q.append(k); rendered.add(k)

    for sec_key in render_q:
        raw_section = secs.get(sec_key, "").strip()
        if not raw_section: continue

        sec_lines    = raw_section.split("\n")
        heading_text = SEC_NAMES.get(sec_key, sec_lines[0].strip().rstrip(":"))
        _add_heading(doc, heading_text, tpl)

        body_lines = sec_lines[1:] if len(sec_lines) > 1 else []
        _render_section_body(doc, body_lines, tpl, accent, fn, fs)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _render_section_body(doc, lines: list, tpl: dict, accent, fn: str, fs: int):
    """
    Render section body with proper formatting:
    - Bullets: List Bullet style, indented, accent colour bullet char
    - Job headers: bold company, italic dates/location
    - Sub-headers (Achievements:): bold, slightly smaller
    - Wrapped PDF lines: rejoined into full paragraphs before rendering
    """
    tc         = RGBColor(0x1A, 0x1A, 0x2E)
    grey       = RGBColor(0x55, 0x55, 0x55)
    BULLET_CHARS = set("•-*–◦·")
    bullet_re  = re.compile(r"^([•\-\*–◦·])\s*(.+)")
    job_pipe_re= re.compile(r"\|")
    date_re    = re.compile(r"\b(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|\d{4}|Present|Current)\b", re.I)
    sub_hdr_re = re.compile(r"(?i)^(achievements?|responsibilities|key\s+achievements?):?\s*$")
    # Lines that are clearly date/location only (short lines with year)
    meta_only_re = re.compile(r"^\(?[A-Z][a-z]+.*\d{4}|\(.*\)|.*\d{4}\s*[–\-]")

    # Step 1: Rejoin lines that were wrapped by PDF extraction
    # A line that doesn't start with bullet/header and the previous line also wasn't
    # a bullet/header and doesn't end in punctuation → likely wrapped
    rejoined: list[str] = []
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        s    = line.strip()
        if not s:
            rejoined.append("")
            i += 1
            continue

        is_bullet    = bool(bullet_re.match(s))
        is_section   = bool(sub_hdr_re.match(s))
        is_job_hdr   = bool(job_pipe_re.search(s)) and (bool(date_re.search(s)) or s.count("|") >= 1)

        if is_bullet:
            # For bullet lines: merge any continuation lines (lines that don't start new block)
            merged = s
            j = i + 1
            while j < len(lines):
                nxt = lines[j].strip()
                if not nxt:
                    break
                nxt_is_bullet  = bool(bullet_re.match(nxt))
                nxt_is_section = bool(sub_hdr_re.match(nxt))
                nxt_is_job_hdr = bool(job_pipe_re.search(nxt)) and (bool(date_re.search(nxt)) or nxt.count("|") >= 1)
                nxt_is_hdr     = bool(re.match(r"(?i)^[A-Z][A-Z\s&/]{3,}$", nxt))
                if nxt_is_bullet or nxt_is_section or nxt_is_job_hdr or nxt_is_hdr:
                    break
                merged = merged + " " + nxt
                j += 1
            rejoined.append(merged)
            i = j
        elif not is_section and not is_job_hdr:
            # Non-bullet, non-section lines: merge continuations
            merged = s
            j = i + 1
            while j < len(lines):
                nxt = lines[j].strip()
                if not nxt:
                    break
                nxt_is_bullet  = bool(bullet_re.match(nxt))
                nxt_is_section = bool(sub_hdr_re.match(nxt))
                nxt_is_job_hdr = bool(job_pipe_re.search(nxt)) and (bool(date_re.search(nxt)) or nxt.count("|") >= 1)
                nxt_is_hdr     = bool(re.match(r"(?i)^[A-Z][A-Z\s&/]{3,}$", nxt))
                if nxt_is_bullet or nxt_is_section or nxt_is_job_hdr or nxt_is_hdr:
                    break
                if merged.endswith((".", "!", "?", ";")):
                    break
                merged = merged + " " + nxt
                j += 1
            rejoined.append(merged)
            i = j
        else:
            rejoined.append(s)
            i += 1

    # Step 2: Render rejoined lines
    # Track whether we are inside a job block (after a job header line)
    # to know that non-bullet text lines are job responsibilities (format as bullets)
    inside_job = False
    POWER_VERBS = {
        "led","managed","developed","designed","built","conducted","applied","spearheaded",
        "standardised","standardized","prepared","mentored","supervised","completed","achieved",
        "upgraded","published","delivered","screened","identified","executed","awarded",
        "co-authored","managed","isolated","authored","advocated","fabricated","engineered",
        "secured","increased","completed","implemented","demonstrated"
    }

    for line in rejoined:
        s = line.strip()
        if not s:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            continue

        m = bullet_re.match(s)
        # Also detect space-indented lines (PDF sometimes uses spaces instead of bullet chars)
        # These look like "  Led Indo-EU..." — leading space + capital verb starting a sentence
        if not m and not sub_hdr_re.match(s) and not (job_pipe_re.search(s) and (date_re.search(s) or s.count("|") >= 1)):
            # Check if original line (before strip) had leading whitespace AND starts with capital
            # AND looks like a job responsibility sentence (not a header)
            if (line.startswith((" ", "\t")) and s and s[0].isupper() 
                and len(s) > 30 and not re.match(r"^[A-Z][A-Z\s&/]{3,}$", s)
                and not s.startswith(("WORK","EDUCATION","SKILLS","CERT","PUBL","PATENT","AWARD","LANG","PROJ"))):
                # Treat as a bullet
                bp = doc.add_paragraph(style="List Bullet")
                bp.paragraph_format.left_indent  = Cm(0.4)
                bp.paragraph_format.space_before = Pt(1)
                bp.paragraph_format.space_after  = Pt(2)
                run = bp.add_run(s)
                _set_run_font(run, fn, fs, color=tc)
                continue

        if m:
            # Bullet point
            content_text = m.group(2).strip()
            while content_text and content_text[0] in BULLET_CHARS:
                content_text = content_text[1:].strip()
            if not content_text:
                continue
            bp = doc.add_paragraph(style="List Bullet")
            bp.paragraph_format.left_indent  = Cm(0.4)
            bp.paragraph_format.space_before = Pt(1)
            bp.paragraph_format.space_after  = Pt(2)
            run = bp.add_run(content_text)
            _set_run_font(run, fn, fs, color=tc)

        elif sub_hdr_re.match(s):
            # "Achievements:" sub-header
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after  = Pt(2)
            run = p.add_run(s)
            _set_run_font(run, fn, fs, bold=True, color=accent)

        elif job_pipe_re.search(s) and (date_re.search(s) or s.count("|") >= 1):
            # Job header line: "Company | Role | Dates"
            inside_job = True  # We are now inside a job block
            parts = [x.strip() for x in s.split("|")]
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after  = Pt(2)
            if parts:
                r1 = p.add_run(parts[0])
                _set_run_font(r1, fn, fs, bold=True, color=accent)
            rest = "  ·  ".join(parts[1:]) if len(parts) > 1 else ""
            if rest:
                r2 = p.add_run("  |  " + rest)
                _set_run_font(r2, fn, fs - 1, italic=True, color=grey)

        elif re.match(r"^\(?[A-Z][a-z]+[,\s].*\d{4}|\(.*\d{4}.*\)$|.*\d{4}\s*[–\-]\s*", s) and len(s) < 80:
            # Date/location only line (short, contains year)
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(2)
            run = p.add_run(s)
            _set_run_font(run, fn, fs - 1, italic=True, color=grey)

        else:
            # Check if this looks like a job responsibility (starts with action verb, >50 chars)
            first_word = s.split()[0].lower().rstrip(".,;:") if s.split() else ""
            is_responsibility = (
                inside_job and
                len(s) > 50 and
                (first_word in POWER_VERBS or
                 (s[0].isupper() and not re.match(r"^[A-Z][A-Z\s&/]{3,}$", s) and
                  not s.startswith(("The ","This ","A ","An ","In ","With ","At ","Our "))))
            )
            if is_responsibility:
                bp = doc.add_paragraph(style="List Bullet")
                bp.paragraph_format.left_indent  = Cm(0.4)
                bp.paragraph_format.space_before = Pt(1)
                bp.paragraph_format.space_after  = Pt(2)
                run = bp.add_run(s)
                _set_run_font(run, fn, fs, color=tc)
            else:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after  = Pt(2)
                run = p.add_run(s)
                _set_run_font(run, fn, fs, color=tc)




def _parse_header_text(text: str) -> dict:
    """Quick regex-based header parser."""
    email_re = re.compile(r"[\w.+-]+@[\w-]+\.[a-z]{2,}", re.I)
    phone_re = re.compile(r"[\+\(]?\d[\d\s\-\(\)\.]{6,}\d")
    linkedin_re = re.compile(r"linkedin\.com/in/[\S]+", re.I)
    github_re = re.compile(r"github\.com/[\S]+", re.I)

    name, email, phone, linkedin, github = "", "", "", "", ""
    for i, line in enumerate(text.split("\n")):
        s = line.strip()
        if not s:
            continue
        if not name:
            name = s
        else:
            if not email:
                m = email_re.search(s)
                if m:
                    email = m.group()
            if not phone:
                m = phone_re.search(s)
                if m:
                    phone = m.group().strip()
            if not linkedin:
                m = linkedin_re.search(s)
                if m:
                    linkedin = m.group()
            if not github:
                m = github_re.search(s)
                if m:
                    github = m.group()

    return {k: v for k, v in [("name", name), ("email", email), ("phone", phone),
                                ("linkedin", linkedin), ("github", github)] if v}


# ── Public API ────────────────────────────────────────────────────────────────

def export_resume_to_docx(resume_data: dict, template_id: str = "classic_professional") -> bytes:
    """
    Export resume to DOCX.
    If resume_data has '_raw' key (from optimizer), build from text.
    Otherwise build from structured form data.
    """
    tpl_id = resume_data.get("template_name") or template_id or "classic_professional"

    # ── Uploaded Word-file template (word_<filename>) ────────────────
    if isinstance(tpl_id, str) and tpl_id.startswith("word_"):
        try:
            from utils.template_manager import load_template_file, apply_word_template
            fname = tpl_id[len("word_"):]
            tpl_bytes = load_template_file(fname)
            if tpl_bytes:
                return apply_word_template(tpl_bytes, resume_data)
        except Exception:
            pass
        tpl_id = "classic_professional"

    if not template_exists(tpl_id):
        tpl_id = "classic_professional"

    raw = resume_data.get("_raw") or resume_data.get("optimized_text")
    if raw and raw.strip():
        parsed_sections = resume_data.get("_parsed_sections")
        return _build_docx_from_text(raw, tpl_id, parsed_sections)
    return _build_docx_from_resume_data(resume_data, tpl_id)


def export_resume_to_txt(resume_data: dict) -> str:
    """Export resume as clean plain text."""
    raw = resume_data.get("_raw") or resume_data.get("optimized_text")
    if raw:
        return raw

    pi = resume_data.get("personal_info", {})
    lines = []

    name = pi.get("name") or resume_data.get("full_name", "")
    if name:
        lines.append(name)
    contact = " | ".join(v for v in [
        pi.get("email"), pi.get("phone"), pi.get("location"),
        pi.get("linkedin"), pi.get("github")
    ] if v)
    if contact:
        lines.append(contact)
    lines.append("")

    if resume_data.get("summary"):
        lines += ["PROFESSIONAL SUMMARY", resume_data["summary"], ""]

    for exp in resume_data.get("experience", []):
        pos = exp.get("position", "")
        com = exp.get("company", "")
        dat = exp.get("dates", "")
        if pos or com:
            lines.append(f"{pos} | {com}" if (pos and com) else (pos or com))
        if dat:
            lines.append(dat)
        desc = exp.get("description", "")
        if desc:
            lines += desc.split("\n")
        lines.append("")

    for edu in resume_data.get("education", []):
        deg = edu.get("degree", "")
        sch = edu.get("school", "")
        yr = str(edu.get("year", ""))
        if deg or sch:
            lines.append(f"{deg} — {sch}" if (deg and sch) else (deg or sch))
        if yr:
            lines.append(yr)
        lines.append("")

    skills = resume_data.get("skills", "")
    if isinstance(skills, list):
        skills = ", ".join(skills)
    if skills:
        lines += ["SKILLS", skills, ""]

    for key in ["certifications", "publications", "patents", "awards"]:
        val = resume_data.get(key)
        if val:
            lines.append(key.upper())
            if isinstance(val, list):
                lines += val
            else:
                lines.append(str(val))
            lines.append("")

    return "\n".join(lines)


def get_template_preview_html(template_id: str) -> str:
    """Return a small HTML preview of the template (used in the UI)."""
    tpl = _resolve_template(template_id)
    accent_hex = f"#{str(tpl['accent'])}"
    font = tpl["font"]
    border_style = f"border-bottom:2px solid {accent_hex}" if tpl["heading_style"] == "border" else f"text-decoration:underline;color:{accent_hex}"

    return f"""
<div style="padding:10px;font-family:'{font}',Arial,sans-serif;font-size:8px;line-height:1.4;color:#222;width:100%">
  <div style="text-align:center;font-size:13px;font-weight:700;color:{accent_hex};margin-bottom:3px">CANDIDATE NAME</div>
  <div style="text-align:center;font-size:7px;color:#666;margin-bottom:8px">email@example.com  |  +91 98765 43210  |  Hyderabad, India</div>
  <div style="font-size:8px;font-weight:700;color:{accent_hex};{border_style};padding-bottom:2px;margin-bottom:4px">PROFESSIONAL SUMMARY</div>
  <div style="font-size:7.5px;color:#333;margin-bottom:6px">Results-driven professional with expertise in data analysis and strategic planning, seeking to leverage skills in a growth-oriented role...</div>
  <div style="font-size:8px;font-weight:700;color:{accent_hex};{border_style};padding-bottom:2px;margin-bottom:4px">WORK EXPERIENCE</div>
  <div style="font-weight:700;font-size:7.5px;color:{accent_hex}">Senior Analyst  |  Acme Corp</div>
  <div style="font-style:italic;color:#888;font-size:7px;margin-bottom:2px">Jan 2022 – Present  |  Hyderabad</div>
  <div style="font-size:7px;color:#333">• Developed Python ML pipelines reducing analysis time by 40%</div>
  <div style="font-size:7px;color:#333">• Led cross-functional team of 8 delivering ₹2Cr revenue impact</div>
  <div style="font-size:8px;font-weight:700;color:{accent_hex};{border_style};padding-bottom:2px;margin:5px 0 3px">EDUCATION</div>
  <div style="font-weight:700;font-size:7.5px;color:{accent_hex}">M.Tech Computer Science  |  IIT Hyderabad</div>
  <div style="font-style:italic;color:#888;font-size:7px">2019 – 2021  |  CGPA: 8.7/10</div>
</div>"""
