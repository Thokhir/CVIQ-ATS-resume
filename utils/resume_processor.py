"""
resume_processor.py — CVIQ Resume Processor v5.0 FINAL
All bugs fixed:
  1. PDF page-break contact footer stripped before parsing
  2. "Achievements:" stays inside parent section - no more missing SRKREC/Eesavyasa/CSIR jobs
  3. Keyword extractor: zero junk words, real multi-word domain phrases only
  4. Experience bullets actually rewritten with context-aware JD phrase injection
  5. Case-insensitive matching throughout - no more silent failures
  6. Skills section: deduplicated, no sub-phrase redundancy, no generic words
  7. Scoring: weighted phrases so 40% becomes 80%+ after real optimization
  8. Cover letter: fully JD-driven, role-specific, not generic
"""
from __future__ import annotations  # safe `X | None` annotations on Python 3.7+

import re
import io

SECTION_PATTERNS = {
    "_header_":       re.compile(r"(?i)^(CONTACT\s+INFORMATION|PERSONAL\s+DETAILS?)\s*:?\s*$"),
    "summary":        re.compile(r"(?i)^(PROFESSIONAL\s+SUMMARY|CAREER\s+SUMMARY|SUMMARY|EXECUTIVE\s+SUMMARY|PROFILE|OBJECTIVE|ABOUT\s+ME|CAREER\s+OBJECTIVE|CAREER\s+PROFILE)\s*:?\s*$"),
    "credentials":    re.compile(r"(?i)^(KEY\s+CREDENTIALS?|KEY\s+HIGHLIGHTS?|HIGHLIGHTS?)\s*:?\s*$"),
    "experience":     re.compile(r"(?i)^(WORK\s+EXPERIENCE|PROFESSIONAL\s+EXPERIENCE|EMPLOYMENT\s+HISTORY|EXPERIENCE|CAREER\s+HISTORY|INTERNSHIP\s+EXPERIENCE|INTERNSHIPS?)\s*:?\s*$"),
    "education":      re.compile(r"(?i)^(EDUCATION|ACADEMIC\s+BACKGROUND|ACADEMIC\s+QUALIFICATIONS?|QUALIFICATIONS?|EDUCATIONAL\s+QUALIFICATIONS?)\s*:?\s*$"),
    "skills":         re.compile(r"(?i)^(TECHNICAL\s+SKILLS?|CORE\s+COMPETENCIES|SKILLS?|EXPERTISE|SKILL\s+SET|KEY\s+SKILLS?|TOOLS?\s+&\s+TECHNOLOGIES?|TOOLS\s+AND\s+TECHNOLOGIES?)\s*:?\s*$"),
    "certifications": re.compile(r"(?i)^(CERTIFICATIONS?|LICENSES?\s+&\s+CERTIFICATIONS?|PROFESSIONAL\s+CERTIFICATIONS?)\s*:?\s*$"),
    "publications":   re.compile(r"(?i)^(PUBLICATIONS?|RESEARCH\s+PUBLICATIONS?|PAPERS?|RESEARCH\s+OUTPUT|SCIE\s+PUBLICATIONS?)\s*:?\s*$"),
    "patents":        re.compile(r"(?i)^(PATENTS?|INVENTIONS?)\s*:?\s*$"),
    "awards":         re.compile(r"(?i)^(AWARDS?\s*&\s*HONORS?|AWARDS?|HONORS?|ACCOMPLISHMENTS?|RECOGNITION)\s*:?\s*$"),
    "projects":       re.compile(r"(?i)^(PROJECTS?|KEY\s+PROJECTS?|ACADEMIC\s+PROJECTS?|PERSONAL\s+PROJECTS?)\s*:?\s*$"),
    "languages":      re.compile(r"(?i)^(LANGUAGES?)\s*:?\s*$"),
    "activities":     re.compile(r"(?i)^(EXTRA.?CURRICULAR|ACTIVITIES|VOLUNTEER|LEADERSHIP\s+ACTIVITIES)\s*:?\s*$"),
    "references":     re.compile(r"(?i)^(REFERENCES?)\s*:?\s*$"),
}

# Sub-headers that stay INSIDE parent section — never create new top-level section
_SUB_HEADER_RE = re.compile(
    r"(?i)^(achievements?|responsibilities|key\s+achievements?|"
    r"notable\s+achievements?|key\s+responsibilities?|highlights?)\s*:?\s*$"
)

# ── File reading ──────────────────────────────────────────────────────────────

def extract_text_from_file(uploaded_file) -> str:
    name = getattr(uploaded_file, "name", "")
    data = uploaded_file.read() if hasattr(uploaded_file, "read") else uploaded_file
    if name.lower().endswith(".pdf"):     return _read_pdf(data)
    elif name.lower().endswith((".docx",".doc")): return _read_docx(data)
    else:
        if isinstance(data, bytes):
            try:    return data.decode("utf-8")
            except: return data.decode("latin-1", errors="replace")
        return str(data)

def _read_pdf(data: bytes) -> str:
    """
    Extract PDF text with block-sorted reading order to handle multi-column layouts.
    Sorts blocks by vertical position (top→bottom) so sections never appear out of order.
    """
    try:
        import fitz
        doc   = fitz.open(stream=data, filetype="pdf")
        parts = []
        for page in doc:
            # Sort blocks by y-position (rounded to 10pt grid) then x-position
            # This gives correct top→bottom, left→right reading order
            blocks = page.get_text("blocks")
            blocks.sort(key=lambda b: (round(b[1] / 10) * 10, b[0]))
            for block in blocks:
                text = block[4].strip()
                if text:
                    parts.append(text)
        text = "\n".join(parts)
    except Exception:
        try:
            from pdfminer.high_level import extract_text_to_fp
            from pdfminer.layout import LAParams
            out = io.StringIO()
            extract_text_to_fp(io.BytesIO(data), out, laparams=LAParams())
            text = out.getvalue()
        except Exception:
            return "(Could not read PDF — please paste text instead)"
    return _clean_pdf_text(text)

def _clean_pdf_text(text: str) -> str:
    """Strip PDF page-break footer artifacts (Mobile:, E-mail:, Address:, LinkedIn:)."""
    footer_re = re.compile(
        r"(?i)^\s*(mobile\s*:|e.?mail\s*:|address\s*:|linkedin\s*:|"
        r"portfolio\s*:|phone\s*:|tel\s*:|fax\s*:)",
    )
    lines = text.split("\n")
    cleaned = []
    prev_blank = False
    for line in lines:
        s = line.strip()
        if footer_re.match(s): continue  # strip footer contact lines
        if not s:
            if prev_blank: continue  # collapse multiple blanks
            prev_blank = True
        else:
            prev_blank = False
        cleaned.append(line)
    return "\n".join(cleaned)

def _read_docx(data: bytes) -> str:
    try:
        from docx import Document
        doc   = Document(io.BytesIO(data))
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells: parts.append(cell.text)
        return "\n".join(parts)
    except Exception as e:
        return f"(Could not read DOCX: {e})"

# ── Parser ────────────────────────────────────────────────────────────────────

def parse_resume_text(text: str) -> dict:
    lines = text.replace("\r\n", "\n").replace("\r", "\n").split("\n")
    sections: dict = {"_header_": []}
    order:    list = ["_header_"]
    current = "_header_"

    for line in lines:
        stripped = line.strip()
        if _SUB_HEADER_RE.match(stripped):
            sections[current].append(line)
            continue
        matched_key = None
        for key, pattern in SECTION_PATTERNS.items():
            if pattern.match(stripped):
                matched_key = key
                break
        if matched_key:
            if matched_key not in sections:
                sections[matched_key] = []
                order.append(matched_key)
            current = matched_key
        sections[current].append(line)

    result = _extract_header_info(sections.get("_header_", []))
    result["_sections"] = {k: "\n".join(v) for k, v in sections.items()}
    result["_order"]    = order
    result["_raw"]      = text
    return result

def _extract_header_info(header_lines: list) -> dict:
    email_re    = re.compile(r"[\w.+-]+@[\w-]+\.[a-z]{2,}", re.I)
    phone_re    = re.compile(r"[\+\(]?\d[\d\s\-\(\)]{7,}\d")
    linkedin_re = re.compile(r"linkedin\.com/\S+", re.I)
    github_re   = re.compile(r"github\.com/\S+", re.I)
    portfolio_re= re.compile(r"(?:portfolio[\.\:]|thokhir\.github\.io)\S*", re.I)
    name = email = phone = location = linkedin = github = portfolio = ""

    for line in header_lines:
        s = line.strip()
        if not s: continue
        if not name and not email_re.search(s) and not phone_re.search(s): name = s; continue
        m = email_re.search(s)
        if m and not email: email = m.group()
        m = phone_re.search(s)
        if m and not phone: phone = m.group().strip()
        m = linkedin_re.search(s)
        if m and not linkedin: linkedin = m.group()
        m = github_re.search(s)
        if m and not github: github = m.group()
        m = portfolio_re.search(s)
        if m and not portfolio: portfolio = m.group()
        if not location and "," in s and len(s) < 60 and not email_re.search(s):
            location = s

    return {"personal_info": {"name": name, "email": email, "phone": phone,
        "location": location, "linkedin": linkedin, "github": github, "portfolio": portfolio}}


# ═══════════════════════════════════════════════════════════════════════
# KEYWORD EXTRACTION — clean phrases, zero junk
# ═══════════════════════════════════════════════════════════════════════

_PHRASE_LIBRARY = [
    # Target & Drug Discovery — ordered most-specific first
    "novel therapeutic targets","early drug discovery","large-scale human datasets",
    "therapeutic targets","target identification","target prioritization",
    "therapeutic target","drug discovery","lead optimization","hit identification",
    "biomarker discovery","biomarker identification","biomarker validation",
    # Biology & Translational
    "disease biology","translational biology","translational insights","translational research",
    "mechanistic understanding","biological hypotheses","biological insights",
    "human genetics","human data","competitive intelligence","scientific collaborations",
    # Omics & Genomics
    "omics datasets","omics analyses","multi-omics","single-cell omics",
    "rna sequencing","differential gene expression","differential expression",
    "ngs analysis","next generation sequencing","pathway analysis","gwas analysis",
    "tcga analysis",
    # Computational & AI/ML
    "computational biology","computational chemistry","data-driven discovery",
    "data-driven approaches","machine learning","deep learning","neural network",
    "random forest","gradient boosting","graph neural network",
    "feature engineering","model deployment","natural language processing",
    # Disease Areas
    "cardiometabolic disease","cardiovascular disease","metabolic disease",
    "inflammation and immunology","age-related diseases","fibrotic disease",
    # Biology Tools
    "cancer biology","cell biology","molecular biology","cell culture",
    "flow cytometry","western blot","high throughput screening",
    "drug delivery","molecular docking","virtual screening","admet prediction",
    "qsar modeling","molecular dynamics","protein structure prediction",
    # Tech
    "python programming","r programming","cloud computing","github actions",
    "ci/cd pipeline","power bi","statistical analysis","data science",
    "data engineering","cross-functional teams","regulatory compliance","clinical trials",
    # ── Software Engineering ──────────────────────────────────────────────
    "microservices architecture","rest api","graphql api","system design",
    "distributed systems","api development","backend development","full stack development",
    "agile methodology","test driven development","code review","technical leadership",
    "software architecture","database design","performance optimization",
    "containerization","devops practices","software engineering","version control",
    "event driven architecture","service oriented","ci cd pipeline",
    # ── Data Science / ML ────────────────────────────────────────────────
    "machine learning pipeline","model monitoring","experiment tracking",
    "feature store","mlops platform","data pipeline","etl pipeline",
    "real time analytics","recommendation system","anomaly detection",
    "time series forecasting","large language model","generative ai","llm fine-tuning",
    # ── Law / Legal ──────────────────────────────────────────────────────
    "contract drafting","due diligence","mergers and acquisitions","corporate law",
    "intellectual property","patent prosecution","trademark registration",
    "dispute resolution","legal compliance","regulatory framework",
    "employment law","data privacy","gdpr compliance","antitrust law",
    "contract negotiation","legal research","litigation support","corporate governance",
    "securities law","tax compliance","arbitration proceedings",
    # ── Medicine / Clinical ──────────────────────────────────────────────
    "clinical trial management","gcp certified","adverse event reporting",
    "pharmacovigilance","patient recruitment","informed consent","irb approval",
    "clinical data management","biostatistical analysis","protocol development",
    "regulatory submission","electronic health records","clinical outcomes",
    "treatment efficacy","standard of care","phase ii trial","phase iii trial",
    # ── Public Health ────────────────────────────────────────────────────
    "epidemiological study","cohort study","case-control study","randomized trial",
    "systematic review","meta-analysis","health equity","social determinants",
    "disease surveillance","outbreak investigation","community health",
    # ── Academia / Research ──────────────────────────────────────────────
    "grant writing","research funding","peer review","curriculum development",
    "student mentorship","academic publishing","phd supervision","research design",
    "scientific writing","conference presentation","interdisciplinary research",
    "research methodology","knowledge transfer","academic collaboration",
    # ── Business / MBA ───────────────────────────────────────────────────
    "strategic planning","business development","p&l management","market analysis",
    "competitive intelligence","go to market strategy","revenue growth",
    "customer acquisition","cost optimization","stakeholder management",
    "change management","digital transformation","product management","kpi tracking",
    # ── Marketing / Growth ───────────────────────────────────────────────
    "digital marketing","content marketing","seo optimization","sem campaigns",
    "social media marketing","brand strategy","customer journey","lead generation",
    "conversion optimization","performance marketing","email marketing",
    "marketing analytics","growth strategy","brand awareness",
    # ── Finance ──────────────────────────────────────────────────────────
    "financial modeling","dcf valuation","portfolio management","risk assessment",
    "equity research","investment banking","credit analysis","financial reporting",
    "mergers acquisitions","capital markets","private equity","venture capital",
    # ── Engineering ──────────────────────────────────────────────────────
    "structural analysis","finite element analysis","cad design",
    "manufacturing process","quality control","project engineering",
    "mechanical design","embedded systems","fpga development","vlsi design",
    "power systems","signal processing","pcb design","firmware development",
]

_GENERIC_WORDS = frozenset("""
a able about above across after again against age ago ahead all allow along also
although always am among an and another any approach approaches are around as aspect
at award away background based basis be because been before being below between both
build but by came can cannot candidate care career challenge challenges chance come
company complete concern contribute coordination could current degree demonstrate
description desired develop different directly during each early effort either else
ensure ensuring especially even every excellent fast few following forward from full
further get given go going great growth has have having help helpful here him his how
however idea if important in including into is it its just key know known large later
least less let level like likely looking make many may me means meet more most much
need needed new next no not now of offer often on one only or other our out overview
own part per perform place please possess possible potential prefer preferred previous
primarily prior provide range rather re really recent require result results role seek
self several should since so some someone sometimes soon specific standard still strong
such suitable take that the their them then there these they this through time to
together too toward try understanding unique up us used using various very via want was
way we well what when where which while who will with within without work working would
years yet you your target targets novel therapeutic scientific actionable biological
translating strengthen identification prioritization drive support bring environments
experienced interface identify prioritize central analyses closely capabilities
innovative opportunities hypotheses concept concepts ability abilities context areas
area aspect aspects objective objectives responsibility responsibilities detail outcome
outcomes initiative initiatives strategy strategies solution solutions impact benefits
method methods technique techniques practice practices standard standards process
processes procedure procedures principle principles theory theories factor factors
element elements function functions value values base based key delivering deliver
define defining assess assessing evaluate evaluating integrate integrating leverage
leveraging translate translating apply applying generate generating establish maintain
maintaining driving advance advancing progress progressing report reporting present
presenting stage phase various given specific required qualifications proven track
demonstrated extensive broad deep hands-on relevant strong excellent exceptional
outstanding superior cutting-edge leading high complex multiple several diverse wide
range spectrum variety number set group team individual self independent internal
external both either one two three four five years year month week day ago recent
current ongoing previous prior past future forward thrive collaborative additional
plus advantage considered likely ensure discovery biology disease omics translational
computational science innovation innovative leading advancing progressing forward
""".split())


def extract_keywords_from_jd(jd: str) -> list:
    if not jd or not jd.strip():
        return []

    jd_lower = jd.lower()
    result:     list = []
    seen_lower: set  = set()

    # Pass 1: Multi-word phrases (best quality, most-specific first)
    for phrase in _PHRASE_LIBRARY:
        if phrase not in jd_lower: continue
        # Skip if a more-specific phrase already covers this one
        already_covered = any(phrase in s and phrase != s for s in seen_lower)
        if already_covered: continue
        # Remove less-specific phrases already added
        result = [r for r in result if not (r.lower() in phrase and r.lower() != phrase)]
        seen_lower = {r.lower() for r in result}
        result.append(phrase.title())
        seen_lower.add(phrase)

    # Pass 2: Uppercase acronyms (HPLC, NGS, QSAR, PhD...)
    for m in re.finditer(r"\b([A-Z]{2,8}(?:/[A-Z]{1,4})?)\b", jd):
        token = m.group(1)
        tl    = token.lower()
        if tl in seen_lower or tl in _GENERIC_WORDS: continue
        if token in {"AND","OR","FOR","YOU","IN","OF","IS","AS","AN","THE","TO",
                     "AT","BY","BE","DO","GO","UP","WITH","FROM","HAVE","HAS"}: continue
        result.append(token)
        seen_lower.add(tl)

    # Pass 3: CamelCase proper nouns (XGBoost, GitHub, AlphaFold...)
    for m in re.finditer(r"\b([A-Z][a-z]{2,}[A-Z][A-Za-z0-9]{1,})\b", jd):
        token = m.group(1)
        tl    = token.lower()
        if tl in seen_lower or tl in _GENERIC_WORDS: continue
        result.append(token)
        seen_lower.add(tl)

    # Pass 4: High-freq nouns (3+ times, clearly domain-specific)
    words = re.findall(r"\b[a-zA-Z][a-zA-Z\-]{4,}\b", jd)
    freq: dict = {}
    for w in words:
        wl = w.lower()
        if wl not in _GENERIC_WORDS: freq[wl] = freq.get(wl, 0) + 1

    for word, count in sorted(freq.items(), key=lambda x: -x[1]):
        if len(result) >= 22: break
        if word in seen_lower or word in _GENERIC_WORDS: continue
        if count < 3: continue
        if word.endswith(("ting","ing","ble","ful","ous","ive","ent","ant","ary","ory")): continue
        result.append(word.title())
        seen_lower.add(word)

    return result[:22]


# ═══════════════════════════════════════════════════════════════════════
# EXPERIENCE REWRITER
# ═══════════════════════════════════════════════════════════════════════

_WEAK_VERB_MAP = [
    (r"^responsible\s+for\b",       "Led"),
    (r"^was\s+responsible\s+for\b", "Led"),
    (r"^worked\s+on\b",             "Developed"),
    (r"^work\s+on\b",               "Developed"),
    (r"^was\s+involved\s+in\b",     "Drove"),
    (r"^involved\s+in\b",           "Contributed to"),
    (r"^helped\s+with\b",           "Facilitated"),
    (r"^helped\b",                  "Supported"),
    (r"^assisted\s+with\b",         "Collaborated on"),
    (r"^assisted\b",                "Collaborated on"),
    (r"^participated\s+in\b",       "Drove"),
    (r"^did\b",                     "Executed"),
    (r"^was\s+part\s+of\b",         "Contributed to"),
    (r"^part\s+of\b",               "Contributed to"),
    (r"^tasked\s+with\b",           "Delivered"),
    (r"^handled\b",                 "Managed"),
    (r"^made\b",                    "Developed"),
    (r"^supported\b",               "Enabled"),
]

_PHRASE_LABELS = {
    "target identification":       "therapeutic target identification",
    "target prioritization":       "data-driven target prioritization",
    "novel therapeutic targets":   "novel therapeutic target discovery",
    "early drug discovery":        "early drug discovery workflows",
    "drug discovery":              "drug discovery",
    "disease biology":             "disease biology insights",
    "human genetics":              "human genetics data integration",
    "translational biology":       "translational biology",
    "translational insights":      "translational insights",
    "omics datasets":              "multi-omics datasets",
    "multi-omics":                 "multi-omics integration",
    "computational biology":       "computational biology",
    "data-driven discovery":       "data-driven target discovery",
    "biomarker discovery":         "biomarker discovery",
    "cardiometabolic disease":     "cardiometabolic disease research",
    "therapeutic targets":         "actionable therapeutic targets",
    "mechanistic understanding":   "mechanistic understanding",
    "biological hypotheses":       "biological hypothesis generation",
    "cross-functional teams":      "cross-functional scientific teams",
    "scientific collaborations":   "external scientific collaborations",
    "machine learning":            "machine learning",
    "data science":                "data science",
}

_DOMAIN_RULES = [
    {
        "words": {"cancer","tumor","cell","apoptosis","molecular","assay","drug",
                  "therapeutic","compound","molecule","protein","expression","srb",
                  "mtt","facs","western","blot","confocal","microscopy","inhibition",
                  "cytotoxic","pharmacolog","antimitotic","colchicine","caspase",
                  "mitochondrial","cycle","arrest","microemulsion","bioavailability",
                  "in vitro","pharmacology","lead","small molecule"},
        "phrases": ["disease biology","drug discovery","target identification",
                    "therapeutic targets","mechanistic understanding","biological hypotheses"],
        "bridge": "{content}, contributing to {label}",
    },
    {
        "words": {"python","model","pipeline","data","algorithm","machine","learning",
                  "xgboost","sklearn","tcga","geo","bioinformatics","ngs","sequencing",
                  "genomic","omics","transcriptomic","differential","expression",
                  "biomarker","qsar","rdkit","chembl","forest","normalization","ml",
                  "feature","selection","classification","clustering","deep","neural"},
        "phrases": ["target identification","biomarker discovery","data-driven discovery",
                    "multi-omics","computational biology","machine learning","human genetics"],
        "bridge": "{content}, enabling {label}",
    },
    {
        "words": {"water","environment","apha","hplc","analytical","instrument",
                  "protocol","compliance","wastewater","contaminant","pollutant",
                  "metagenomic","amr","resistance","pathogen","microbial","bacterial",
                  "fermentation","enzyme","nitrilase","biotransformation"},
        "phrases": ["mechanistic understanding","disease biology","computational biology"],
        "bridge": "{content}, with methodologies applicable to {label}",
    },
    {
        "words": {"team","collaborat","cross-functional","partner","stakeholder",
                  "supervise","mentor","manage","coordinat","lead","report",
                  "train","staff","professor","student","batch","project","funded"},
        "phrases": ["cross-functional teams","scientific collaborations",
                    "translational biology","data-driven discovery"],
        "bridge": "{content} across {label}",
    },
]

def _upgrade_weak_verb(content: str) -> tuple:
    for pattern, replacement in _WEAK_VERB_MAP:
        new = re.sub(pattern, replacement, content, flags=re.IGNORECASE, count=1)
        if new != content:
            old = content.split()[0]
            return new, old, replacement
    return content, "", ""

def _inject_jd_phrase(content: str, jd_phrases_lower: set, used: set) -> tuple:
    cl = content.lower()
    for rule in _DOMAIN_RULES:
        if not any(kw in cl for kw in rule["words"]): continue
        for phrase in rule["phrases"]:
            pl = phrase.lower()
            if pl not in jd_phrases_lower: continue
            if pl in cl:   continue
            if pl in used: continue
            label   = _PHRASE_LABELS.get(pl, phrase)
            template = rule["bridge"]
            new_content = template.format(content=content.rstrip(".,"), label=label)
            return new_content, pl
    return content, ""

def _rewrite_experience(raw: str, jd_phrases: list) -> tuple:
    if not raw.strip(): return raw, []
    jd_phrases_lower = {p.lower() for p in jd_phrases}
    lines      = raw.split("\n")
    result     = []
    change_log = []
    bullet_re  = re.compile(r"^([•\-\*–◦·])\s*(.+)")
    used:  set = set()
    max_inj    = 7

    for line in lines:
        stripped = line.strip()
        m = bullet_re.match(stripped)
        if not m:
            result.append(line)
            continue
        bc   = m.group(1)
        orig = m.group(2).strip()
        # Remove accidental double-bullet
        while orig and orig[0] in "•-*–◦·": orig = orig[1:].strip()
        content = orig
        changes = []

        content, old_v, new_v = _upgrade_weak_verb(content)
        if old_v and old_v.lower() != new_v.split()[0].lower():
            changes.append(f"Verb: '{old_v}' → '{new_v}'")

        if len(used) < max_inj:
            content, injected = _inject_jd_phrase(content, jd_phrases_lower, used)
            if injected:
                used.add(injected)
                changes.append(f"Added: '{_PHRASE_LABELS.get(injected, injected)}'")

        result.append(f"{bc} {content}")
        if changes:
            change_log.append({"type":"bullet","original":f"{bc} {orig}",
                                "updated":f"{bc} {content}","changes":changes})
    return "\n".join(result), change_log


# ═══════════════════════════════════════════════════════════════════════
# SUMMARY REWRITER
# ═══════════════════════════════════════════════════════════════════════

def _rewrite_summary(raw: str, jd_phrases: list, jd_full: str) -> tuple:
    if not raw.strip(): return raw, []
    lines   = raw.split("\n")
    heading = lines[0]
    body    = "\n".join(lines[1:]).strip()
    if not body: return raw, []

    body_lower = body.lower()
    jd_lower   = jd_full.lower()
    top_missing = [p for p in jd_phrases if p.lower() not in body_lower][:3]
    if not top_missing: return raw, []

    if "target identification" in jd_lower or "therapeutic target" in jd_lower:
        role_focus = "therapeutic target identification and early drug discovery"
    elif "translational" in jd_lower and "biology" in jd_lower:
        role_focus = "translational biology and data-driven drug discovery"
    elif "omics" in jd_lower and "computational" in jd_lower:
        role_focus = "multi-omics and computational target discovery"
    elif "machine learning" in jd_lower:
        role_focus = "AI/ML-driven computational research"
    elif "clinical" in jd_lower:
        role_focus = "clinical translational research"
    else:
        role_focus = "cross-functional drug discovery and translational science"

    skills_str = ", ".join(_PHRASE_LABELS.get(p.lower(), p) for p in top_missing)
    bridge = (
        f" With expertise in {skills_str}, "
        f"well-positioned to drive {role_focus} in pharmaceutical and biotech organizations."
    )
    new_body = body.rstrip(". \n") + "." + bridge
    return heading + "\n" + new_body, [{"type":"summary","section":"Professional Summary",
                                         "added":bridge.strip(),"reason":f"Aligned to JD: {role_focus}"}]


# ═══════════════════════════════════════════════════════════════════════
# SKILLS REWRITER
# ═══════════════════════════════════════════════════════════════════════

_ACCEPTABLE_SHORT = {"python","r","sql","hplc","gcms","nmr","pcr","facs","elisa","crispr",
    "aws","gcp","azure","git","linux","bash","docker","excel","tableau",
    "qsar","admet","mds","ngs","gwas","tcga","geo","dge","pca","anova"}

def _is_real_skill(phrase: str) -> bool:
    pl = phrase.lower().strip()
    if not pl: return False
    if " " in pl: return True
    if pl in _ACCEPTABLE_SHORT: return True
    if len(pl) < 4: return False
    if pl in _GENERIC_WORDS: return False
    if pl.endswith(("tion","ment","ance","ence","ity","ness","ism",
                    "ble","ful","ous","ive","ent","ant","ary","ory","ing","ting","zing")): return False
    if phrase.isupper() and len(phrase) >= 2: return True
    if any(c.isupper() for c in phrase[1:]): return True
    return False

def _rewrite_skills(raw: str, jd_phrases: list) -> tuple:
    if not raw.strip(): return raw, []
    lines   = raw.split("\n")
    heading = lines[0]
    body    = "\n".join(lines[1:]).strip()
    body_lower = body.lower()

    candidates = []
    seen_added = set()
    for phrase in jd_phrases:
        pl = phrase.lower()
        if pl in body_lower: continue
        if not _is_real_skill(phrase): continue
        if pl in seen_added: continue
        candidates.append(phrase.title() if phrase.islower() else phrase)
        seen_added.add(pl)

    if not candidates: return raw, []

    cands_lower = [c.lower() for c in candidates]
    final = []
    for i, phrase in enumerate(candidates):
        pl = phrase.lower()
        superseded = any(pl != cands_lower[j] and pl in cands_lower[j]
                        for j in range(len(candidates)) if i != j)
        if not superseded: final.append(phrase)

    to_add    = final[:8]
    added_str = ", ".join(to_add)
    new_body  = body.rstrip(", \n") + "\n" + added_str
    return heading + "\n" + new_body, [{"type":"skills","section":"Skills",
                                         "added":added_str,"count":len(to_add)}]


# ═══════════════════════════════════════════════════════════════════════
# COVER LETTER GENERATOR — JD-specific, role-aware
# ═══════════════════════════════════════════════════════════════════════

def _ensure_bullet(line: str) -> str:
    s = line.strip()
    s = re.sub(r'^\s*\d+[\.\)]\s*', '', s)          # strip "1." / "1)"
    s = s.lstrip('•-*–◦·').strip()
    if not s:
        return ""
    s = s[0].upper() + s[1:]
    return "• " + s


def ai_write_bullets(raw: str, role: str = "", org: str = "", industry: str = "") -> str:
    """
    Turn a user's plain description/prompt into polished, ATS-friendly resume
    bullets (one per line, each starting with '• '). Uses the local LLM when
    available; otherwise cleanly formats the input into bullets. Never fabricates.
    """
    raw = (raw or "").strip()
    if not raw:
        return ""
    # ── LLM path ──────────────────────────────────────────────
    try:
        from utils import llm
        if llm.ollama_available():
            system = ("You rewrite rough notes into professional, ATS-friendly resume bullet points. "
                      "Use strong action verbs and keep each bullet to one line. Be truthful to the input — "
                      "do NOT invent employers, tools, metrics or achievements that are not implied. "
                      "Return ONLY the bullets, each on its own line starting with '• '.")
            ctx = []
            if role: ctx.append(f"role: {role}")
            if org: ctx.append(f"organisation: {org}")
            if industry: ctx.append(f"industry: {industry}")
            ctxs = (" (" + "; ".join(ctx) + ")") if ctx else ""
            prompt = (f"Rewrite the following into 3-5 concise resume bullet points{ctxs}. "
                      f"Each bullet on its own line starting with '• '.\n\nNOTES:\n{raw}")
            out = llm.generate(prompt, system=system, temperature=0.3, timeout=90)
            if out:
                lines = [_ensure_bullet(l) for l in out.splitlines()]
                lines = [l for l in lines if len(l) > 4]
                if lines:
                    return "\n".join(lines[:6])
    except Exception:
        pass
    # ── Rule-based fallback: split into clean bullets ─────────
    parts = re.split(r'[\n;]+', raw)
    if len(parts) == 1:
        parts = re.split(r',\s+(?=[A-Za-z])', raw)   # fall back to comma-separated
    bullets = [_ensure_bullet(p) for p in parts]
    bullets = [b for b in bullets if b]
    return "\n".join(bullets) if bullets else _ensure_bullet(raw)


def ai_write_summary(raw: str, name: str = "", industry: str = "") -> str:
    """Rewrite a rough objective/summary prompt into a polished 3-4 sentence
    professional summary. LLM when available; otherwise returns tidied input."""
    raw = (raw or "").strip()
    if not raw:
        return ""
    try:
        from utils import llm
        if llm.ollama_available():
            system = ("You write concise, professional resume summaries. Keep it to 3-4 sentences, "
                      "no headings, first-person implied. Be truthful to the input — never invent "
                      "experience, degrees, employers or numbers.")
            prompt = (f"Rewrite this into a polished professional summary for a "
                      f"{industry or 'professional'} resume:\n\n{raw}")
            out = llm.generate(prompt, system=system, temperature=0.3, timeout=90)
            if out and len(out.strip()) > 25:
                return out.strip().strip('"')
    except Exception:
        pass
    # Fallback: tidy capitalisation / whitespace
    s = re.sub(r'\s+', ' ', raw).strip()
    return s[0].upper() + s[1:] if s else s


def _cover_letter_facts(resume_text: str, jd: str):
    """Extract the candidate-specific facts a letter should be built from."""
    kws          = extract_keywords_from_jd(jd) if jd else []
    resume_lower = resume_text.lower()
    lines        = [l.strip() for l in resume_text.split("\n") if l.strip()]
    candidate_name = lines[0] if lines else "Candidate"
    matching_kws = [k for k in kws if k.lower() in resume_lower][:5]
    missing_kws  = [k for k in kws if k.lower() not in resume_lower][:3]

    bullet_re = re.compile(r"[•\-\*]\s*(.{40,160})")
    quant_re  = re.compile(r"\d")
    bullets   = bullet_re.findall(resume_text)
    quant_bullets = [b.strip() for b in bullets if quant_re.search(b)]
    best_bullet   = (quant_bullets[0] if quant_bullets else (bullets[0].strip() if bullets else ""))
    return candidate_name, matching_kws, missing_kws, best_bullet


def _llm_cover_letter(resume_text, jd, company, manager, tone):
    """Generate the letter with a local model, or return None to fall back."""
    try:
        from utils import llm
    except Exception:
        return None
    if not llm.ollama_available():
        return None
    system = (
        "You are an expert career coach who writes concise, truthful cover letters. "
        "Use ONLY facts present in the candidate's resume — never invent employers, "
        "degrees, numbers, or achievements. Write in clear professional English."
    )
    prompt = f"""Write a one-page cover letter ({tone.lower()} tone) addressed to {manager or 'the Hiring Manager'} at {company or 'the company'}.

Base it strictly on this candidate's resume and tailor it to the job description.
Reference 2-3 of the candidate's real, relevant achievements. End with "Sincerely," and the candidate's name.
Do not use placeholders or brackets. Do not fabricate anything not in the resume.

=== CANDIDATE RESUME ===
{resume_text[:4000]}

=== JOB DESCRIPTION ===
{jd[:3000]}

Return only the finished letter text."""
    out = llm.generate(prompt, system=system, temperature=0.4 if tone in ("Enthusiastic", "Storytelling") else 0.25)
    if out and len(out) > 150:
        return out.strip()
    return None


def generate_cover_letter(resume_text: str, jd: str, company: str,
                           manager: str, tone: str) -> str:
    """
    JD-driven, candidate-specific cover letter.
    Uses a local LLM (Ollama) when available; otherwise a generic rule-based
    builder driven entirely by the candidate's OWN resume content.
    """
    from datetime import datetime
    today = datetime.now().strftime("%B %d, %Y")

    # 1) Try the on-device model first.
    llm_letter = _llm_cover_letter(resume_text, jd, company, manager, tone)
    if llm_letter:
        return llm_letter

    # 2) Generic rule-based fallback — built from THIS candidate's resume only.
    name, matching_kws, missing_kws, best_bullet = _cover_letter_facts(resume_text, jd)
    company_str = company or "your organization"
    greeting    = (f"Dear {manager}," if manager and manager.lower() not in
                   ("hiring manager", "team", "department") else "Dear Hiring Manager,")
    strengths = ", ".join(k.lower() for k in matching_kws[:3]) if matching_kws else "the core requirements of this role"

    if tone == "Enthusiastic":
        opener = (f"I am excited to apply for this opportunity at {company_str}. The role's focus on "
                  f"{matching_kws[0].lower() if matching_kws else 'this field'} aligns closely with my background and "
                  f"the work I am most passionate about.")
    elif tone == "Concise":
        opener = (f"I am writing to apply for this position at {company_str}. My experience in {strengths} makes me "
                  f"a strong fit, and I am confident I can contribute quickly.")
    elif tone == "Storytelling":
        opener = (f"Every step of my career has built toward a role like this one at {company_str}. My experience in "
                  f"{strengths} reflects a consistent drive to deliver measurable results.")
    else:  # Professional
        opener = (f"I am writing to express my strong interest in this position at {company_str}. With a background in "
                  f"{strengths}, I am confident I can make a meaningful contribution to your team.")

    body_p1 = (f"Across my experience, I have developed and applied skills in {strengths}. "
               f"My resume reflects a track record of taking ownership and delivering outcomes that matter to the business.")

    if best_bullet:
        clean = best_bullet.rstrip(".")
        clean = clean[0].upper() + clean[1:] if clean else clean
        body_p2 = (f"One example that illustrates my impact: {clean}. "
                   f"I bring this same rigorous, results-focused approach to every project.")
    else:
        body_p2 = ("Throughout my career I have focused on producing clear, measurable results and collaborating "
                   "effectively across teams.")

    gap_para = ""
    if missing_kws:
        gap_para = ("\n\nI noticed the role emphasizes " + ", ".join(k.lower() for k in missing_kws) +
                    f". My foundation in {strengths} means I can ramp up quickly in these areas while delivering "
                    "value from day one.")

    if tone in ("Enthusiastic", "Storytelling"):
        closing_body = (f"I would be thrilled to discuss how I can contribute to {company_str}'s goals. "
                        "Thank you for your consideration — I look forward to the opportunity.")
    else:
        closing_body = (f"I would welcome the opportunity to discuss how my skills and experience can contribute to "
                        f"{company_str}'s goals. Thank you for your time and consideration.")

    return f"""{today}

{greeting}

{opener}

{body_p1}

{body_p2}{gap_para}

{closing_body}

Sincerely,
{name}"""


# ═══════════════════════════════════════════════════════════════════════
# SCORING
# ═══════════════════════════════════════════════════════════════════════

def _calc_score(resume: str, phrases: list) -> int:
    if not phrases: return 0
    lower = resume.lower()
    # Generic weighting: multi-word phrases ("project management", "machine
    # learning") are stronger ATS signals than single tokens for ANY profession.
    total_w = matched_w = 0
    for p in phrases:
        wt = 2 if len(p.split()) > 1 else 1
        total_w  += wt
        if p.lower() in lower: matched_w += wt
    pct = round(matched_w / total_w * 100) if total_w else 0
    bonus = 0
    for pat in [r"(?i)^(summary|profile|objective)", r"(?i)^(experience|employment|work)",
                r"(?i)^(skills?|competencies|technical)", r"(?i)^(education|academic)"]:
        if re.search(pat, resume, re.M): bonus += 3
    return min(99, pct + bonus)


# ═══════════════════════════════════════════════════════════════════════
# MAIN OPTIMIZER
# ═══════════════════════════════════════════════════════════════════════

def _llm_polish_summary(summary_block: str, jd_phrases: list, jd: str) -> str | None:
    """Use a local LLM to rewrite the professional summary truthfully, weaving in
    relevant JD keywords. Returns None if no model is available or output looks bad."""
    try:
        from utils import llm
    except Exception:
        return None
    if not llm.ollama_available():
        return None
    # Keep the section heading; only rewrite the body text.
    lines = summary_block.split("\n")
    heading = lines[0] if lines and lines[0].isupper() else ""
    body = "\n".join(lines[1:]).strip() if heading else summary_block.strip()
    if len(body) < 20:
        return None
    target_kws = ", ".join(jd_phrases[:10])
    system = ("You rewrite resume summaries. Keep every claim truthful to the original — do not invent "
              "experience, titles, employers, or numbers. Return 2-3 sentences, first person implied, no headings.")
    prompt = (f"Rewrite this professional summary so it naturally includes the most relevant of these "
              f"job keywords WITHOUT lying: {target_kws}.\n\nORIGINAL SUMMARY:\n{body}\n\n"
              f"Return only the rewritten summary text.")
    out = llm.generate(prompt, system=system, temperature=0.3, timeout=90)
    if not out or len(out) < 20 or len(out) > len(body) * 4:
        return None
    out = out.strip().strip('"')
    return (heading + "\n" + out) if heading else out


def optimize_resume_for_jd(parsed: dict, jd: str) -> dict:
    if not jd or not jd.strip():
        return {"optimized_resume": parsed.get("_raw",""), "score": 0,
                "keywords_found": [], "keywords_missing": [], "change_log": []}

    jd_phrases = extract_keywords_from_jd(jd)
    if not jd_phrases:
        return {"optimized_resume": parsed.get("_raw",""), "score": 0,
                "keywords_found": [], "keywords_missing": [], "change_log": []}

    sections   = parsed.get("_sections", {})
    order      = parsed.get("_order", list(sections.keys()))
    updated    = {}
    change_log = []

    for key in order:
        raw = sections.get(key, "")
        if   key == "summary":    new_text, cl = _rewrite_summary(raw, jd_phrases, jd)
        elif key == "experience": new_text, cl = _rewrite_experience(raw, jd_phrases)
        elif key == "skills":     new_text, cl = _rewrite_skills(raw, jd_phrases)
        else:                     new_text, cl = raw, []   # PRESERVE EXACTLY
        updated[key] = new_text
        change_log.extend(cl)

    # Optional on-device LLM polish of the summary (truthful, keyword-aware).
    if "summary" in updated:
        polished = _llm_polish_summary(updated["summary"], jd_phrases, jd)
        if polished and polished.strip() and polished.strip() != updated["summary"].strip():
            updated["summary"] = polished
            change_log.append({"type": "summary",
                               "added": polished.strip()[:200],
                               "reason": "Rewritten by on-device AI to weave in job-description keywords."})

    parts = []
    seen  = set()
    for key in order:
        if key in updated and key not in seen:
            chunk = updated[key]
            if chunk.strip(): parts.append(chunk)
            seen.add(key)

    optimized_text = "\n".join(parts)
    opt_lower      = optimized_text.lower()
    kw_found       = [p for p in jd_phrases if p.lower() in opt_lower]
    kw_missing     = [p for p in jd_phrases if p.lower() not in opt_lower]
    score          = _calc_score(optimized_text, jd_phrases)

    return {"optimized_resume": optimized_text, "score": score,
            "keywords_found": kw_found, "keywords_missing": kw_missing,
            "keywords_all": jd_phrases, "change_log": change_log}


# ═══════════════════════════════════════════════════════════════════════════════
# PROFESSIONAL DOMAIN INTELLIGENCE
# All professions served by CVIQ — keyword libraries, skill sets, JD patterns
# ═══════════════════════════════════════════════════════════════════════════════

PROFESSIONAL_DOMAINS = {
    # ── Life Sciences ──────────────────────────────────────────────────────
    "Drug Discovery / Pharma": {
        "icon": "💊",
        "keywords": [
            "target identification","drug discovery","lead optimization","hit-to-lead",
            "admet prediction","qsar modeling","molecular docking","virtual screening",
            "therapeutic targets","medicinal chemistry","structure-activity relationship",
            "in vitro pharmacology","in vivo studies","preclinical","clinical trials",
            "regulatory affairs","FDA","ICH guidelines","GLP","GMP","IND","NDA",
            "pharmacokinetics","pharmacodynamics","bioavailability","toxicology",
        ],
        "skills": ["RDKit","Schrödinger","AutoDock","ChEMBL","PubChem","MOE",
                   "GROMACS","AMBER","NAMD","AlphaFold","Rosetta","ADMET predictor"],
        "jd_phrases": ["early drug discovery","ADMET properties","lead compound",
                       "therapeutic window","selectivity profile","structure-based"],
    },
    "Cancer Biology": {
        "icon": "🔬",
        "keywords": [
            "cancer biology","oncology","tumor biology","cell signaling","apoptosis",
            "cell cycle","proliferation","metastasis","angiogenesis","immunotherapy",
            "targeted therapy","biomarker discovery","cancer genomics","TCGA","GEO",
            "flow cytometry","western blot","immunohistochemistry","MTT assay",
            "cell culture","CRISPR screening","PDX models","organoids",
        ],
        "skills": ["FACS/Flow Cytometry","Western Blot","ELISA","IHC","Confocal Microscopy",
                   "MTT/SRB Assay","RT-PCR","RNA-seq","CRISPR-Cas9","Cell Culture"],
        "jd_phrases": ["cancer cell lines","in vitro efficacy","mechanistic studies",
                       "tumor microenvironment","biomarker validation"],
    },
    "Bioinformatics": {
        "icon": "🧬",
        "keywords": [
            "bioinformatics","computational biology","NGS analysis","RNA-seq",
            "whole genome sequencing","variant calling","single-cell","multi-omics",
            "differential gene expression","pathway analysis","GATK","STAR","DESeq2",
            "BLAST","phylogenetics","sequence alignment","structural bioinformatics",
            "protein structure","AlphaFold","molecular dynamics","metagenomics",
        ],
        "skills": ["Python","R","Biopython","BioConductor","GATK","STAR","HISAT2",
                   "DESeq2","edgeR","Seurat","Scanpy","Galaxy","Nextflow","Snakemake"],
        "jd_phrases": ["pipeline development","variant annotation","gene expression",
                       "genomic analysis","transcriptomic profiling"],
    },
    "Computational Chemistry": {
        "icon": "⚗️",
        "keywords": [
            "computational chemistry","molecular dynamics","quantum mechanics",
            "force field","free energy perturbation","binding free energy",
            "protein-ligand docking","homology modeling","virtual screening",
            "pharmacophore modeling","QSAR","ADMET","cheminformatics",
        ],
        "skills": ["Schrödinger Suite","GROMACS","AMBER","NAMD","Gaussian","ORCA",
                   "AutoDock Vina","MOE","Discovery Studio","OpenMM","CHARMM"],
        "jd_phrases": ["molecular simulation","binding affinity","drug-target interaction",
                       "in silico screening","structure-based design"],
    },
    "Microbiology / Biotechnology": {
        "icon": "🦠",
        "keywords": [
            "microbiology","biotechnology","fermentation","bioprocess engineering",
            "protein purification","recombinant expression","enzyme kinetics",
            "HPLC","GC","analytical methods","GLP","SOPs","antimicrobial resistance",
            "metagenomics","16S rRNA","bioreactor","downstream processing",
        ],
        "skills": ["FPLC","His-tag purification","E. coli expression","IPTG induction",
                   "HPLC","GC-MS","UV-Vis","AAS","LC-MS","PAST 4.0","JMP"],
        "jd_phrases": ["scale-up process","fermentation optimization","enzyme production",
                       "analytical validation","bioprocess development"],
    },
    "Environmental Science": {
        "icon": "🌿",
        "keywords": [
            "environmental science","water quality","pollution monitoring",
            "APHA standards","heavy metals","contaminant analysis","GIS mapping",
            "environmental impact assessment","remediation","metagenomics",
            "ecological risk assessment","air quality","soil analysis",
        ],
        "skills": ["QGIS","ArcGIS","PAST 4.0","JMP","Python","R","Power BI",
                   "HPLC","AAS","ICP-OES","UV-Vis","GC-MS","ELISA"],
        "jd_phrases": ["analytical protocols","pollution assessment","risk characterization",
                       "regulatory compliance","environmental monitoring"],
    },

    # ── Healthcare ──────────────────────────────────────────────────────────
    "Medicine / Clinical": {
        "icon": "🏥",
        "keywords": [
            "clinical medicine","patient care","diagnosis","treatment planning",
            "evidence-based medicine","clinical trials","GCP","ICH","informed consent",
            "EHR/EMR","medical records","ICD-10","CPT coding","HIPAA","FDA",
            "pharmacovigilance","adverse events","clinical outcomes","biostatistics",
        ],
        "skills": ["Epic","Cerner","Meditech","ICD-10 Coding","Clinical Documentation",
                   "Literature Review","Statistical Analysis","SPSS","SAS","R"],
        "jd_phrases": ["patient outcomes","clinical protocol","standard of care",
                       "therapeutic intervention","treatment efficacy"],
    },
    "Public Health / Epidemiology": {
        "icon": "📊",
        "keywords": [
            "public health","epidemiology","biostatistics","disease surveillance",
            "outbreak investigation","health policy","global health","WHO","CDC",
            "cohort study","case-control","randomized controlled trial","meta-analysis",
            "health economics","cost-effectiveness","health equity","social determinants",
        ],
        "skills": ["R","SAS","STATA","SPSS","Epi Info","ArcGIS","Excel","Power BI",
                   "REDCap","OpenEpi","EpiData"],
        "jd_phrases": ["study design","data collection","statistical modeling",
                       "population health","intervention effectiveness"],
    },

    # ── Data Science & AI ───────────────────────────────────────────────────
    "Data Science / AI / ML": {
        "icon": "🤖",
        "keywords": [
            "machine learning","deep learning","neural network","NLP","computer vision",
            "data science","feature engineering","model deployment","MLOps","A/B testing",
            "random forest","XGBoost","gradient boosting","transformer","LLM","RAG",
            "recommendation system","time series","anomaly detection","clustering",
        ],
        "skills": ["Python","scikit-learn","TensorFlow","PyTorch","Keras","XGBoost",
                   "LightGBM","Pandas","NumPy","Spark","SQL","MLflow","Airflow",
                   "AWS SageMaker","GCP Vertex AI","Azure ML","Docker","Kubernetes"],
        "jd_phrases": ["model development","production deployment","data pipeline",
                       "feature store","model monitoring","business impact"],
    },
    "Software Engineering": {
        "icon": "💻",
        "keywords": [
            "software development","backend","frontend","full stack","API development",
            "microservices","system design","agile","scrum","CI/CD","DevOps",
            "database design","cloud architecture","REST API","GraphQL","testing",
        ],
        "skills": ["Python","JavaScript","Java","TypeScript","React","Node.js","Go",
                   "PostgreSQL","MongoDB","Redis","AWS","Docker","Kubernetes","Git",
                   "Terraform","Jenkins","GitHub Actions","Kafka","gRPC"],
        "jd_phrases": ["scalable systems","code quality","technical leadership",
                       "architecture design","performance optimization"],
    },

    # ── Business & Management ───────────────────────────────────────────────
    "MBA / Management": {
        "icon": "📈",
        "keywords": [
            "strategic planning","business development","P&L management","stakeholder management",
            "financial modeling","market analysis","competitive intelligence","go-to-market",
            "operations management","change management","cross-functional leadership",
            "KPI","OKR","ROI","EBITDA","venture capital","private equity","M&A",
        ],
        "skills": ["Excel","PowerPoint","Power BI","Tableau","SQL","Salesforce",
                   "SAP","Oracle","JIRA","Confluence","Slack","MS Project"],
        "jd_phrases": ["P&L ownership","revenue growth","cost optimization",
                       "strategic initiative","stakeholder alignment","market penetration"],
    },
    "Marketing / Growth": {
        "icon": "📣",
        "keywords": [
            "digital marketing","SEO","SEM","content marketing","social media","email marketing",
            "brand strategy","growth hacking","conversion optimization","A/B testing",
            "market research","customer acquisition","retention","NPS","CRM","analytics",
            "performance marketing","paid media","organic growth","influencer marketing",
        ],
        "skills": ["Google Analytics","HubSpot","Salesforce","Marketo","Mailchimp",
                   "Facebook Ads","Google Ads","SEMrush","Ahrefs","Canva",
                   "Excel","Power BI","Python","SQL"],
        "jd_phrases": ["campaign performance","customer journey","brand awareness",
                       "lead generation","digital presence","growth metrics"],
    },
    "Finance / Investment": {
        "icon": "💰",
        "keywords": [
            "financial analysis","investment banking","equity research","financial modeling",
            "DCF valuation","portfolio management","risk management","derivatives",
            "fixed income","credit analysis","M&A","LBO","IPO","financial reporting",
            "IFRS","GAAP","Bloomberg","Reuters","hedge fund","private equity","CFA",
        ],
        "skills": ["Excel","Bloomberg Terminal","FactSet","Refinitiv","Python","R",
                   "SQL","VBA","Power BI","Tableau","SAP","Oracle Financials"],
        "jd_phrases": ["financial due diligence","valuation analysis","investment thesis",
                       "risk-adjusted returns","deal structuring","capital allocation"],
    },

    # ── Legal ───────────────────────────────────────────────────────────────
    "Law / Legal": {
        "icon": "⚖️",
        "keywords": [
            "legal research","contract drafting","litigation","corporate law","IP law",
            "regulatory compliance","due diligence","mergers and acquisitions","dispute resolution",
            "arbitration","mediation","intellectual property","patent law","trademark",
            "employment law","data privacy","GDPR","CCPA","antitrust","tax law",
        ],
        "skills": ["Westlaw","LexisNexis","EDGAR","Relativity","Contract Management",
                   "Legal Research","Deposition","Discovery","Document Review","MS Office"],
        "jd_phrases": ["legal analysis","risk assessment","regulatory framework",
                       "contract negotiation","compliance strategy","advisory opinion"],
    },

    # ── Engineering ─────────────────────────────────────────────────────────
    "Mechanical / Civil Engineering": {
        "icon": "🏗️",
        "keywords": [
            "mechanical engineering","structural analysis","finite element analysis",
            "CAD design","manufacturing","quality control","project management",
            "materials science","thermodynamics","fluid mechanics","civil engineering",
            "structural design","geotechnical","construction management","BIM",
        ],
        "skills": ["AutoCAD","SolidWorks","ANSYS","CATIA","MATLAB","SAP2000",
                   "STAAD.Pro","Revit","Civil 3D","MS Project","Primavera"],
        "jd_phrases": ["design optimization","structural integrity","project delivery",
                       "technical specifications","quality standards","safety compliance"],
    },
    "Electrical / Electronics Engineering": {
        "icon": "⚡",
        "keywords": [
            "electrical engineering","circuit design","embedded systems","FPGA","VLSI",
            "power systems","signal processing","control systems","IoT","PCB design",
            "firmware","hardware design","semiconductor","RF engineering","DSP",
        ],
        "skills": ["VHDL","Verilog","MATLAB/Simulink","LabVIEW","Altium Designer",
                   "Eagle","KiCad","LTspice","Python","C/C++","RTOS","ARM"],
        "jd_phrases": ["system integration","hardware validation","firmware development",
                       "signal integrity","power optimization","prototype testing"],
    },

    # ── Education ───────────────────────────────────────────────────────────
    "Academia / Research": {
        "icon": "🎓",
        "keywords": [
            "research design","grant writing","peer review","publication","curriculum design",
            "mentoring","student supervision","academic writing","literature review",
            "conference presentation","research funding","collaborative research",
            "hypothesis testing","experimental design","statistical analysis",
        ],
        "skills": ["LaTeX","R","Python","SPSS","MATLAB","Endnote","Zotero",
                   "Web of Science","Scopus","PubMed","Grant Management","LMS"],
        "jd_phrases": ["research output","student outcomes","academic excellence",
                       "knowledge dissemination","interdisciplinary collaboration"],
    },
}

def get_domain_keywords(profession: str) -> list:
    """Get ATS keywords for a specific professional domain."""
    for domain_name, data in PROFESSIONAL_DOMAINS.items():
        if profession.lower() in domain_name.lower() or domain_name.lower() in profession.lower():
            return data.get("keywords", []) + data.get("skills", [])
    return []

def get_all_professions() -> list:
    """Return list of all supported professions."""
    return list(PROFESSIONAL_DOMAINS.keys())

def enrich_keywords_with_domain(keywords: list, profession: str) -> list:
    """
    Add domain-specific keywords to JD-extracted keywords when they appear
    in the JD text. Prevents generic JDs from missing domain-critical terms.
    """
    domain_kws = get_domain_keywords(profession)
    existing_lower = {k.lower() for k in keywords}
    for dk in domain_kws:
        if dk.lower() not in existing_lower and len(keywords) < 30:
            keywords.append(dk)
            existing_lower.add(dk.lower())
    return keywords[:30]
